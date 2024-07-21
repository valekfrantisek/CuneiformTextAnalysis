from docx import Document

version = 'BETA 0.0.1'
authors = ['František Válek']
project_name = 'CuneiformTextAnalysis'
project_web = 'https://github.com/valekfrantisek/CuneiformTextAnalysis'


def load_document(input_path):
    """ This function loads the document for further analysis. """
    manuscript_document = Document(input_path)
    return manuscript_document


def sum_line_designation(section:str, line_num:str):
    return f'{section} {line_num}'


def assign_type(formatted_data:dict):
    """ This function assigns a type of analysed text based on the concise formatting. """
    text = formatted_data['text']
    if formatted_data['italic']:
        return {'text': text, 'type': 'syllabic'}
    elif formatted_data['superscript'] and ':' in text:
        return {'text': text, 'type': 'manuscript designation'}
    elif formatted_data['small caps']:
        if formatted_data['superscript']:
            return {'text': text, 'type': 'post/determinative'}
        if formatted_data['subscript']:
            return {'text': text, 'type': 'sign number'}
        else:
            return {'text': text, 'type': 'ideo/logographic'}
    elif formatted_data['uppercase']:
        return {'text': text, 'type': 'sign_value'}
    else:
        # print('Check this one', text)
        return {'text': text, 'type': 'other'}


def get_formatted_text_from_cell(input_cell_data):
    """ This function collects data on formatting from the line data from the formatted docx document."""
    formatted_data = [] 
    for paragraph in input_cell_data.paragraphs:
        for run in paragraph.runs:
            text = run.text
            italic = run.italic
            uppercase = text.isupper()
            superscript = run.font.superscript
            subscript = run.font.subscript
            small_caps = run.font.small_caps
            formatted_data.append({
                'text': text,
                'italic': italic,
                'uppercase': uppercase,
                'superscript': superscript,
                'subscript': subscript,
                'small caps': small_caps
            })
    return formatted_data


def simple_print_line_data(formatted_line_data:list):
    """ This function simply prints line data without further information. """
    line_text = ''
    for element in formatted_line_data:
        line_text += element['text']
    
    return line_text


def get_formatted_text_from_table(input_table, section:str, layout='lines'):
    full_formated_data_of_one_table = {'cuneiform_data': {}, 'translation_data': {}}
    full_line_designation = sum_line_designation(section=section, line_num='0')
    next_cell_translation = False
    next_cell_note = False
    next_cell_cuneiform = False
    note_id = 0
    
    if layout == 'lines':
        for row in input_table.rows:
            for i, cell in enumerate(row.cells):
                if i == 0:
                    if cell.text == '':
                        """ If the first cell is empty it is line for translation. """
                        next_cell_translation = True
                        next_cell_note = False
                        next_cell_cuneiform = False
                        
                    elif cell.text.lower() == 'note':
                        """ If the first cell is 'note' it is line for note. """
                        next_cell_note = True
                        next_cell_translation = False
                        next_cell_cuneiform = False
                    else:
                        next_cell_cuneiform = True
                        next_cell_note = False
                        next_cell_translation = False
                        line_num = cell.text
                        full_line_designation = sum_line_designation(section=section, line_num=line_num)
                    
                elif i == 1:
                    if next_cell_translation:
                        full_formated_data_of_one_table['translation_data'][full_line_designation] = cell.text
                        next_cell_translation = False

                    elif next_cell_cuneiform:
                        cell_data = get_formatted_text_from_cell(input_cell_data=cell)
                        line_primary_analysis = []
                        for element in cell_data:
                            element_info = assign_type(element)
                            line_primary_analysis.append(element_info)
                            
                        full_formated_data_of_one_table['cuneiform_data'][full_line_designation] = line_primary_analysis
                        next_cell_cuneiform = False
                        
                    elif next_cell_note:
                        full_formated_data_of_one_table['cuneiform_data'][f'note {full_line_designation}-{note_id}'] = cell.text
                        full_formated_data_of_one_table['translation_data'][f'note {full_line_designation}-{note_id}'] = cell.text
                        note_id += 1
                        next_cell_note = False
                                
        return full_formated_data_of_one_table
    
    elif layout == 'columns':
        for row in input_table.rows:
            for i, cell in enumerate(row.cells):
                if i == 0:
                    if cell.text == '':
                        """ If the first cell is empty it is an error. """
                        next_cell_note = False
                        next_cell_cuneiform = False
                        print('Error, lines must be designated!')
                        
                    elif cell.text.lower() == 'note':
                        """ If the first cell is 'note' it is line for note. """
                        next_cell_note = True
                        next_cell_cuneiform = False
                    else:
                        next_cell_cuneiform = True
                        next_cell_note = False
                        line_num = cell.text
                        full_line_designation = sum_line_designation(section=section, line_num=line_num)
                    
                elif i == 1:
                    if next_cell_cuneiform:
                        cell_data = get_formatted_text_from_cell(input_cell_data=cell)
                        line_primary_analysis = []
                        for element in cell_data:
                            element_info = assign_type(element)
                            line_primary_analysis.append(element_info)
                            
                        full_formated_data_of_one_table['cuneiform_data'][full_line_designation] = line_primary_analysis
                        
                    elif next_cell_note:
                        full_formated_data_of_one_table['cuneiform_data'][f'note {full_line_designation}-{note_id}'] = cell.text
                        full_formated_data_of_one_table['translation_data'][f'note {full_line_designation}-{note_id}'] = cell.text
                        note_id += 1
                    
                elif i == 2 and not next_cell_note:
                    """ The second column is for translation. """
                    full_formated_data_of_one_table['translation_data'][full_line_designation] = cell.text
                                
        return full_formated_data_of_one_table
    
    elif layout == 'no_translation':
        for row in input_table.rows:
            for i, cell in enumerate(row.cells):
                if i == 0:
                    if cell.text == '':
                        """ If the first cell is empty it is an error. """
                        next_cell_note = False
                        next_cell_cuneiform = False
                        print('Error, lines must be designated!')
                        
                    elif cell.text.lower() == 'note':
                        """ If the first cell is 'note' it is line for note. """
                        next_cell_note = True
                        next_cell_cuneiform = False
                    else:
                        next_cell_cuneiform = True
                        next_cell_note = False
                        line_num = cell.text
                        full_line_designation = sum_line_designation(section=section, line_num=line_num)
                    
                elif i == 1:
                    if next_cell_cuneiform:
                        cell_data = get_formatted_text_from_cell(input_cell_data=cell)
                        line_primary_analysis = []
                        for element in cell_data:
                            element_info = assign_type(element)
                            line_primary_analysis.append(element_info)
                            
                        full_formated_data_of_one_table['cuneiform_data'][full_line_designation] = line_primary_analysis
                        full_formated_data_of_one_table['translation_data'][full_line_designation] = '---add translation---'
                        
                    elif next_cell_note:
                        full_formated_data_of_one_table['cuneiform_data'][f'note {full_line_designation}-{note_id}'] = cell.text
                        full_formated_data_of_one_table['translation_data'][f'note {full_line_designation}-{note_id}'] = cell.text
                        note_id += 1
                                
        return full_formated_data_of_one_table
    
    
def get_headings_and_tables_from_document(input_document):
    """ This function gets headings and tables from DOCX document. """
    document_data = []
    table_idx = 0
    for paragraph in input_document.paragraphs:
        if paragraph.style.name.startswith('Heading'):
            heading_text = paragraph.text
            # Find next table
            while table_idx < len(input_document.tables):
                next_elem = paragraph._element.getnext()
                if next_elem is not None and next_elem.tag.endswith('tbl'):
                    table = input_document.tables[table_idx]
                    document_data.append({
                        'heading': heading_text,
                        'table': table
                    })
                    table_idx += 1
                    break
                table_idx += 1
                
    return document_data


def extract_data_from_composition_document(input_path, layout):
    """ This function analyses document that contains transcriptions of cuneiform texts. """
    # TODO: provide detail description of this function
    # TODO: add functionality that ignores non-heading and non-tabular data
    # TODO: add functionality that works without tables

    input_document = Document(input_path)
    
    document_headings_et_tables = get_headings_and_tables_from_document(input_document=input_document)
    
    complete_analysed_document = {}
    
    for section in document_headings_et_tables:
        heading = section['heading']
        table = section['table']
    
        full_formated_data_of_one_table = get_formatted_text_from_table(input_table=table, section=heading, layout=layout)
        
        complete_analysed_document[heading] = full_formated_data_of_one_table
        
        # NOTE: printing individual lines in simple format
        for line in full_formated_data_of_one_table['cuneiform_data']:
            if 'note' in line:
                print(full_formated_data_of_one_table['cuneiform_data'][line])
            else:
                # print(full_formated_data_of_one_table[line])
                line_text = simple_print_line_data(full_formated_data_of_one_table['cuneiform_data'][line])
                print(line, line_text, full_formated_data_of_one_table['translation_data'][line])
                
    return complete_analysed_document